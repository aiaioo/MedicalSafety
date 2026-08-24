import base64
import io
import json
import re
import uuid
from datetime import datetime, timezone
from html import escape as html_escape
from html.parser import HTMLParser
from pathlib import Path
from urllib.parse import parse_qs, urlsplit

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.image.image import Image as DocxImage
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from flask import Flask, Response, abort, jsonify, render_template, request, send_from_directory, url_for

from converters import ConversionError, convert_to_pdf

BASE_DIR = Path(__file__).resolve().parent
DOCUMENTS_DIR = BASE_DIR / "documents"
STORAGE_DIR = BASE_DIR / "storage"
CACHE_DIR = STORAGE_DIR / "cache"
ANNOTATIONS_DIR = STORAGE_DIR / "annotations"
SNIPPETS_DIR = STORAGE_DIR / "snippets"
REPORTS_DIR = STORAGE_DIR / "reports"

for d in (DOCUMENTS_DIR, CACHE_DIR, ANNOTATIONS_DIR, SNIPPETS_DIR, REPORTS_DIR):
    d.mkdir(parents=True, exist_ok=True)

DOC_ID_RE = re.compile(r"^[A-Za-z0-9_-]+$")
HEX_COLOR_RE = re.compile(r"^#[0-9a-fA-F]{6}$")
DEFAULT_ANNOTATION_COLOR = "#e02424"

# Report page margins, in points — left/right/header/footer. Defaults match the
# geometry render_report_pdf has always used (fitz mediabox inset of
# (36, 46, -36, -46)pt), so existing reports render unchanged until a user
# explicitly opens Page setup and changes them.
REPORT_DEFAULT_MARGINS = {"left": 36, "right": 36, "header": 46, "footer": 46}
REPORT_MARGIN_MIN = 0
REPORT_MARGIN_MAX = 200  # pt; keeps left+right and header+footer well under A4's 595x842pt


def sanitize_margins(raw, fallback=None):
    fallback = fallback if isinstance(fallback, dict) else REPORT_DEFAULT_MARGINS
    result = {}
    for key, default in REPORT_DEFAULT_MARGINS.items():
        val = raw.get(key) if isinstance(raw, dict) else None
        try:
            val = float(val)
        except (TypeError, ValueError):
            val = None
        if val is None or not (REPORT_MARGIN_MIN <= val <= REPORT_MARGIN_MAX):
            val = fallback.get(key, default)
        result[key] = val
    return result

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB, generous for scanned case files

UPLOAD_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".doc": "docx"}


class DocumentError(Exception):
    def __init__(self, message, status=404):
        super().__init__(message)
        self.message = message
        self.status = status


@app.errorhandler(DocumentError)
def handle_document_error(err):
    if request.path.startswith("/api/"):
        return jsonify({"error": err.message}), err.status
    return render_template("error.html", message=err.message), err.status


# ---------------------------------------------------------------------------
# Document resolution helpers
# ---------------------------------------------------------------------------

def normalize_type(raw_type):
    t = (raw_type or "pdf").strip().lower()
    if t == "pdf":
        return "pdf"
    if t in ("docx", "doc", "word"):
        return "docx"
    raise DocumentError(f"Unsupported document type: {raw_type!r}", 400)


def check_doc_id(doc_id):
    if not doc_id or not DOC_ID_RE.match(doc_id):
        raise DocumentError(f"Invalid document id: {doc_id!r}", 400)


def check_report_id(report_id):
    if not report_id or not DOC_ID_RE.match(report_id):
        raise DocumentError(f"Invalid document id: {report_id!r}", 400)


def resolve_pdf_path(doc_id, raw_type):
    check_doc_id(doc_id)
    norm_type = normalize_type(raw_type)

    if norm_type == "pdf":
        path = DOCUMENTS_DIR / f"{doc_id}.pdf"
        if not path.exists():
            raise DocumentError(f"No PDF found for doc '{doc_id}' (expected {path.name} in documents/)")
        return path, norm_type

    src = None
    for ext in (".docx", ".doc"):
        candidate = DOCUMENTS_DIR / f"{doc_id}{ext}"
        if candidate.exists():
            src = candidate
            break
    if src is None:
        raise DocumentError(f"No Word document found for doc '{doc_id}' (expected .docx or .doc in documents/)")

    cached = CACHE_DIR / f"{doc_id}.pdf"
    if not cached.exists() or src.stat().st_mtime > cached.stat().st_mtime:
        try:
            converted = convert_to_pdf(src, CACHE_DIR)
        except ConversionError as exc:
            raise DocumentError(str(exc), 500) from exc
        if converted != cached:
            converted.replace(cached)
    return cached, norm_type


# ---------------------------------------------------------------------------
# JSON storage helpers
# ---------------------------------------------------------------------------

def load_json(path: Path, default):
    if not path.exists():
        return default
    with open(path) as f:
        return json.load(f)


def save_json(path: Path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(f".{uuid.uuid4().hex}.tmp")
    with open(tmp, "w") as f:
        json.dump(data, f, indent=2)
    tmp.replace(path)


def annotations_path(doc_id, norm_type):
    return ANNOTATIONS_DIR / f"{doc_id}__{norm_type}.json"


def snippets_meta_path(doc_id, norm_type):
    return SNIPPETS_DIR / f"{doc_id}__{norm_type}.json"


def snippets_dir(doc_id, norm_type):
    return SNIPPETS_DIR / f"{doc_id}__{norm_type}"


def report_path(report_id):
    return REPORTS_DIR / f"{report_id}.json"


def slugify_report_name(name):
    slug = re.sub(r"[^a-z0-9]+", "-", name.strip().lower()).strip("-")
    return slug[:50] or "document"


# ---------------------------------------------------------------------------
# Baking annotations into a snippet render
# ---------------------------------------------------------------------------

def _clamp01(value):
    return max(0.0, min(1.0, float(value)))


def _hex_to_rgb01(hex_color):
    hex_color = hex_color.lstrip("#")
    return (
        int(hex_color[0:2], 16) / 255,
        int(hex_color[2:4], 16) / 255,
        int(hex_color[4:6], 16) / 255,
    )


def sanitize_annotations(raw, max_count=500, max_points=2000):
    """Validate/clamp client-submitted annotations before drawing them into a PDF page."""
    if not isinstance(raw, list):
        return []
    out = []
    for item in raw[:max_count]:
        if not isinstance(item, dict):
            continue
        color = item.get("color")
        if not (isinstance(color, str) and HEX_COLOR_RE.match(color)):
            color = DEFAULT_ANNOTATION_COLOR

        if item.get("kind") == "rect":
            try:
                x, y, w, h = (_clamp01(item[k]) for k in ("x", "y", "w", "h"))
            except (KeyError, TypeError, ValueError):
                continue
            if w <= 0 or h <= 0:
                continue
            out.append({"kind": "rect", "color": color, "x": x, "y": y, "w": w, "h": h})

        elif item.get("kind") == "freehand":
            pts = item.get("points")
            if not isinstance(pts, list):
                continue
            clean_pts = []
            for pt in pts[:max_points]:
                if not (isinstance(pt, list) and len(pt) == 2):
                    continue
                try:
                    clean_pts.append((_clamp01(pt[0]), _clamp01(pt[1])))
                except (TypeError, ValueError):
                    continue
            if len(clean_pts) >= 2:
                out.append({"kind": "freehand", "color": color, "points": clean_pts})
    return out


def draw_annotations_on_page(page, annotations, page_rect):
    """Burn sanitized annotations into a fitz page's content stream (in memory only,
    never saved back to disk) so a subsequent get_pixmap() render includes them
    as sharp vector shapes at whatever DPI is requested."""
    if not annotations:
        return
    shape = page.new_shape()
    line_width = max(1.2, page_rect.width * 0.0025)
    for a in annotations:
        color = _hex_to_rgb01(a["color"])
        if a["kind"] == "rect":
            r = fitz.Rect(
                page_rect.x0 + a["x"] * page_rect.width,
                page_rect.y0 + a["y"] * page_rect.height,
                page_rect.x0 + (a["x"] + a["w"]) * page_rect.width,
                page_rect.y0 + (a["y"] + a["h"]) * page_rect.height,
            )
            shape.draw_rect(r)
            shape.finish(color=color, width=line_width)
        elif a["kind"] == "freehand":
            pts = [
                fitz.Point(page_rect.x0 + px * page_rect.width, page_rect.y0 + py * page_rect.height)
                for px, py in a["points"]
            ]
            shape.draw_polyline(pts)
            shape.finish(color=color, width=line_width, closePath=False)
    shape.commit()


# ---------------------------------------------------------------------------
# Report editor: HTML sanitization, image inlining, PDF rendering
# ---------------------------------------------------------------------------

REPORT_ALLOWED_TAGS = {
    "p", "br", "div", "span", "h1", "h2", "h3", "h4",
    "strong", "b", "em", "i", "u", "s",
    "ul", "ol", "li", "a", "img", "blockquote", "hr",
    "table", "thead", "tbody", "tr", "td", "th",
    "figure", "figcaption", "pre", "code",
}
REPORT_VOID_TAGS = {"br", "img", "hr"}
REPORT_STRIP_CONTENT_TAGS = {"script", "style", "iframe", "object", "embed", "form", "input", "button", "svg"}
REPORT_SAFE_URL_RE = re.compile(
    r"^(https?://|mailto:|/media/|data:image/(png|jpeg|jpg|gif|webp);base64,)", re.IGNORECASE
)
REPORT_SAFE_STYLE_RE = re.compile(r"^[a-zA-Z0-9\s:;#%.,\-()]*$")


def _clean_report_attrs(tag, attrs):
    out = {}
    for name, value in attrs:
        if value is None:
            continue
        name = name.lower()
        if name == "class":
            out["class"] = value
        elif name == "style":
            low = value.lower()
            if REPORT_SAFE_STYLE_RE.match(value) and "expression" not in low and "javascript" not in low:
                out["style"] = value
        elif tag == "a" and name == "href":
            if REPORT_SAFE_URL_RE.match(value.strip()):
                out["href"] = value
        elif tag == "img" and name == "src":
            if REPORT_SAFE_URL_RE.match(value.strip()):
                out["src"] = value
        elif tag == "img" and name in ("alt", "width", "height"):
            out[name] = value
        elif tag in ("td", "th") and name in ("colspan", "rowspan") and value.isdigit():
            out[name] = value
    return out


class _ReportHTMLSanitizer(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.out = []
        self.skip_depth = 0

    def _attr_str(self, tag, attrs):
        clean = _clean_report_attrs(tag, attrs)
        return "".join(f' {k}="{html_escape(v, quote=True)}"' for k, v in clean.items())

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in REPORT_STRIP_CONTENT_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth or tag not in REPORT_ALLOWED_TAGS:
            return
        self.out.append(f"<{tag}{self._attr_str(tag, attrs)}>")

    def handle_startendtag(self, tag, attrs):
        tag = tag.lower()
        if self.skip_depth or tag in REPORT_STRIP_CONTENT_TAGS or tag not in REPORT_ALLOWED_TAGS:
            return
        self.out.append(f"<{tag}{self._attr_str(tag, attrs)}/>")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in REPORT_STRIP_CONTENT_TAGS:
            if self.skip_depth:
                self.skip_depth -= 1
            return
        if self.skip_depth or tag not in REPORT_ALLOWED_TAGS or tag in REPORT_VOID_TAGS:
            return
        self.out.append(f"</{tag}>")

    def handle_data(self, data):
        if not self.skip_depth:
            self.out.append(html_escape(data))

    def get_html(self):
        return "".join(self.out)


def sanitize_report_html(raw_html, max_len=500_000):
    """Whitelist-sanitize client-submitted rich text before it is stored or
    ever re-rendered (in a browser or fed into the PDF Story renderer)."""
    if not isinstance(raw_html, str):
        return ""
    parser = _ReportHTMLSanitizer()
    parser.feed(raw_html[:max_len])
    parser.close()
    return parser.get_html()


REPORT_IMG_SRC_RE = re.compile(r'src="(/media/snippets/[^"]*)"')


def inline_report_images(html_str):
    """Replace <img src="/media/snippets/..."> references with base64 data URIs
    read directly from the snippet files on disk, so the exported PDF is
    self-contained and Story (which has no network access) can render it."""

    def repl(match):
        url = match.group(1)
        parsed = urlsplit(url)
        parts = [p for p in parsed.path.split("/") if p]
        if len(parts) < 2:
            return match.group(0)
        filename, url_doc_id = parts[-1], parts[-2]
        if not DOC_ID_RE.match(url_doc_id) or "/" in filename or "\\" in filename:
            return match.group(0)
        try:
            url_type = normalize_type(parse_qs(parsed.query).get("type", ["pdf"])[0])
        except DocumentError:
            return match.group(0)
        f = snippets_dir(url_doc_id, url_type) / filename
        if not f.exists():
            return match.group(0)
        b64 = base64.b64encode(f.read_bytes()).decode("ascii")
        return f'src="data:image/png;base64,{b64}"'

    return REPORT_IMG_SRC_RE.sub(repl, html_str)


REPORT_PDF_CSS = """
  body { font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #1c1c1c; }
  h1 { font-size: 20pt; margin-bottom: 4pt; }
  h2 { font-size: 15pt; }
  h3 { font-size: 13pt; }
  img { max-width: 100%; display: block; margin: 12pt 0; }
  figure { margin: 12pt 0; }
  figure img { margin: 0; }
  figcaption { font-size: 9pt; color: #555; }
  table { border-collapse: collapse; width: 100%; }
  td, th { border: 1px solid #ccc; padding: 4pt; text-align: left; }
"""


def render_report_pdf(title, body_html, margins=None):
    heading = f"<h1>{html_escape(title)}</h1>" if title else ""
    full_html = f"<html><head><style>{REPORT_PDF_CSS}</style></head><body>{heading}{body_html}</body></html>"

    m = sanitize_margins(margins)
    mediabox = fitz.paper_rect("a4")
    where = mediabox + (m["left"], m["header"], -m["right"], -m["footer"])
    story = fitz.Story(html=full_html)
    buf = io.BytesIO()
    writer = fitz.DocumentWriter(buf)
    more = 1
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
    writer.close()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word (.docx) export
#
# python-docx has no HTML importer, so we walk the same whitelisted tag set
# (REPORT_ALLOWED_TAGS) that sanitize_report_html/render_report_pdf already
# rely on and build the document directly, tag by tag.
# ---------------------------------------------------------------------------

class _HTMLNode:
    __slots__ = ("tag", "attrs", "children")

    def __init__(self, tag, attrs=None):
        self.tag = tag
        self.attrs = attrs or {}
        self.children = []


class _HTMLTreeBuilder(HTMLParser):
    """Lenient HTML -> tree parser. Auto-closes mismatched tags by popping
    the stack up to the nearest matching ancestor, so it tolerates the same
    imperfect nesting a browser's contenteditable might occasionally emit."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _HTMLNode("root")
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        node = _HTMLNode(tag, dict(attrs))
        self.stack[-1].children.append(node)
        if tag not in REPORT_VOID_TAGS:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        self.stack[-1].children.append(_HTMLNode(tag, dict(attrs)))

    def handle_endtag(self, tag):
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                return

    def handle_data(self, data):
        if data:
            self.stack[-1].children.append(data)


DOCX_BLOCK_TAGS = {
    "p", "div", "h1", "h2", "h3", "h4", "ul", "ol", "li",
    "blockquote", "table", "figure", "figcaption", "pre", "hr",
}
DOCX_HEADING_STYLE = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3", "h4": "Heading 4"}


def _docx_image_stream(src):
    if not src or not src.startswith("data:image/"):
        return None
    try:
        header, b64 = src.split(",", 1)
        return io.BytesIO(base64.b64decode(b64))
    except (ValueError, base64.binascii.Error):
        return None


def _docx_add_hyperlink(paragraph, url, text, fmt):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run_el = paragraph._p.makeelement(qn("w:r"), {})
    rpr = paragraph._p.makeelement(qn("w:rPr"), {})
    color = paragraph._p.makeelement(qn("w:color"), {qn("w:val"): "1155CC"})
    underline = paragraph._p.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rpr.append(color)
    rpr.append(underline)
    if fmt.get("bold"):
        rpr.append(paragraph._p.makeelement(qn("w:b"), {}))
    if fmt.get("italic"):
        rpr.append(paragraph._p.makeelement(qn("w:i"), {}))
    run_el.append(rpr)
    text_el = paragraph._p.makeelement(qn("w:t"), {})
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _docx_apply_run_format(run, fmt):
    run.bold = bool(fmt.get("bold"))
    run.italic = bool(fmt.get("italic"))
    run.underline = bool(fmt.get("underline"))
    run.font.strike = bool(fmt.get("strike"))
    if fmt.get("code"):
        run.font.name = "Courier New"


def _docx_add_image(node, paragraph, max_width_emu):
    stream = _docx_image_stream(node.attrs.get("src"))
    if stream is None:
        return
    try:
        info = DocxImage.from_blob(stream.getvalue())
    except Exception:
        return
    native_w = info.width
    width = Emu(min(int(native_w), max_width_emu)) if max_width_emu else None
    stream.seek(0)
    run = paragraph.add_run()
    if width:
        run.add_picture(stream, width=width)
    else:
        run.add_picture(stream)


def _docx_render_inline(node, paragraph, fmt, max_width_emu):
    if isinstance(node, str):
        if node:
            run = paragraph.add_run(node)
            _docx_apply_run_format(run, fmt)
        return

    if node.tag == "br":
        paragraph.add_run().add_break()
        return
    if node.tag == "img":
        _docx_add_image(node, paragraph, max_width_emu)
        return

    child_fmt = dict(fmt)
    if node.tag in ("strong", "b"):
        child_fmt["bold"] = True
    elif node.tag in ("em", "i"):
        child_fmt["italic"] = True
    elif node.tag == "u":
        child_fmt["underline"] = True
    elif node.tag == "s":
        child_fmt["strike"] = True
    elif node.tag == "code":
        child_fmt["code"] = True

    if node.tag == "a":
        href = node.attrs.get("href")
        text = "".join(c for c in node.children if isinstance(c, str))
        if href and text:
            _docx_add_hyperlink(paragraph, href, text, child_fmt)
            return

    for child in node.children:
        _docx_render_inline(child, paragraph, child_fmt, max_width_emu)


def _docx_render_blocks(nodes, doc, max_width_emu, list_ctx=None):
    for node in nodes:
        if isinstance(node, str):
            if node.strip():
                p = doc.add_paragraph()
                p.add_run(node)
            continue

        tag = node.tag
        if tag in DOCX_HEADING_STYLE:
            p = doc.add_paragraph(style=DOCX_HEADING_STYLE[tag])
            for child in node.children:
                _docx_render_inline(child, p, {}, max_width_emu)
        elif tag == "hr":
            p = doc.add_paragraph()
            p.add_run("—" * 20)
        elif tag in ("p", "div"):
            p = doc.add_paragraph()
            for child in node.children:
                _docx_render_inline(child, p, {}, max_width_emu)
        elif tag == "blockquote":
            p = doc.add_paragraph(style="Intense Quote")
            for child in node.children:
                _docx_render_inline(child, p, {}, max_width_emu)
        elif tag in ("ul", "ol"):
            _docx_render_blocks(node.children, doc, max_width_emu, list_ctx=tag)
        elif tag == "li":
            style = "List Number" if list_ctx == "ol" else "List Bullet"
            p = doc.add_paragraph(style=style)
            for child in node.children:
                if isinstance(child, _HTMLNode) and child.tag in DOCX_BLOCK_TAGS:
                    _docx_render_blocks([child], doc, max_width_emu)
                else:
                    _docx_render_inline(child, p, {}, max_width_emu)
        elif tag == "pre":
            text = _flatten_text(node)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
        elif tag == "img":
            p = doc.add_paragraph()
            _docx_render_inline(node, p, {}, max_width_emu)
        elif tag == "figure":
            for child in node.children:
                if isinstance(child, _HTMLNode) and child.tag in DOCX_BLOCK_TAGS:
                    _docx_render_blocks([child], doc, max_width_emu)
                else:
                    p = doc.add_paragraph()
                    _docx_render_inline(child, p, {}, max_width_emu)
        elif tag == "figcaption":
            p = doc.add_paragraph()
            run = p.add_run("".join(_flatten_text(c) if isinstance(c, _HTMLNode) else c for c in node.children))
            run.italic = True
            run.font.size = Pt(9)
        elif tag == "table":
            _docx_render_table(node, doc, max_width_emu)
        else:
            # Unknown/structural wrapper: descend into its children.
            _docx_render_blocks(node.children, doc, max_width_emu, list_ctx=list_ctx)


def _flatten_text(node):
    if isinstance(node, str):
        return node
    return "".join(_flatten_text(c) for c in node.children)


def _docx_render_table(table_node, doc, max_width_emu):
    rows = []
    for section in table_node.children:
        if isinstance(section, _HTMLNode) and section.tag in ("thead", "tbody"):
            rows.extend(c for c in section.children if isinstance(c, _HTMLNode) and c.tag == "tr")
        elif isinstance(section, _HTMLNode) and section.tag == "tr":
            rows.append(section)
    if not rows:
        return
    n_cols = max((sum(1 for c in r.children if isinstance(c, _HTMLNode) and c.tag in ("td", "th")) for r in rows), default=0)
    if n_cols == 0:
        return
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    for r in rows:
        cells = [c for c in r.children if isinstance(c, _HTMLNode) and c.tag in ("td", "th")]
        row_cells = table.add_row().cells
        for i, cell_node in enumerate(cells[:n_cols]):
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            is_header = cell_node.tag == "th"
            for child in cell_node.children:
                if isinstance(child, _HTMLNode) and child.tag in DOCX_BLOCK_TAGS:
                    for gc in child.children:
                        _docx_render_inline(gc, p, {"bold": is_header}, max_width_emu)
                else:
                    _docx_render_inline(child, p, {"bold": is_header}, max_width_emu)


def render_report_docx(title, body_html, margins=None):
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    section = doc.sections[0]
    m = sanitize_margins(margins)
    section.left_margin = Pt(m["left"])
    section.right_margin = Pt(m["right"])
    section.top_margin = Pt(m["header"])
    section.bottom_margin = Pt(m["footer"])
    max_width_emu = int(section.page_width - section.left_margin - section.right_margin)

    if title:
        doc.add_paragraph(title, style="Heading 1")

    builder = _HTMLTreeBuilder()
    builder.feed(body_html)
    builder.close()
    _docx_render_blocks(builder.root.children, doc, max_width_emu)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Pages
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    pdf_ids = {f.stem for f in DOCUMENTS_DIR.glob("*.pdf")}
    docx_ids = {f.stem for f in DOCUMENTS_DIR.glob("*.docx")} | {f.stem for f in DOCUMENTS_DIR.glob("*.doc")}
    docs = [{"id": i, "type": "pdf"} for i in sorted(pdf_ids)]
    docs += [{"id": i, "type": "docx"} for i in sorted(docx_ids - pdf_ids)]
    docs.sort(key=lambda d: d["id"])
    return render_template("index.html", docs=docs)


@app.route("/annotations")
def page_view():
    doc_id = request.args.get("doc", "")
    raw_type = request.args.get("type", "pdf")
    try:
        page = int(request.args.get("page", 1))
    except ValueError:
        raise DocumentError("page must be an integer", 400)

    path, norm_type = resolve_pdf_path(doc_id, raw_type)
    with fitz.open(path) as d:
        page_count = d.page_count
    if page_count == 0:
        raise DocumentError("Document has no pages", 400)
    page = max(1, min(page, page_count))

    return render_template(
        "annotations.html",
        doc_id=doc_id,
        doc_type=raw_type,
        norm_type=norm_type,
        page=page,
        page_count=page_count,
    )


@app.route("/document")
def document_view():
    pdf_ids = {f.stem for f in DOCUMENTS_DIR.glob("*.pdf")}
    docx_ids = {f.stem for f in DOCUMENTS_DIR.glob("*.docx")} | {f.stem for f in DOCUMENTS_DIR.glob("*.doc")}
    source_docs = [{"id": i, "type": "pdf"} for i in sorted(pdf_ids)]
    source_docs += [{"id": i, "type": "docx"} for i in sorted(docx_ids - pdf_ids)]
    source_docs.sort(key=lambda d: d["id"])

    report_id = request.args.get("report", "")
    if report_id:
        check_report_id(report_id)
        if not report_path(report_id).exists():
            raise DocumentError(f"No document with id {report_id!r}", 404)

    preselect_source = request.args.get("source", "")
    preselect_type = request.args.get("type", "pdf")
    if preselect_source:
        check_doc_id(preselect_source)

    return render_template(
        "document.html",
        source_docs=source_docs,
        report_id=report_id,
        preselect_source=preselect_source,
        preselect_type=preselect_type,
    )


# ---------------------------------------------------------------------------
# API: document upload
# ---------------------------------------------------------------------------

@app.route("/api/documents/upload", methods=["POST"])
def api_upload_document():
    f = request.files.get("file")
    if f is None or not f.filename:
        raise DocumentError("No file uploaded (expected multipart field 'file')", 400)

    ext = Path(f.filename).suffix.lower()
    norm_type = UPLOAD_EXTENSIONS.get(ext)
    if norm_type is None:
        raise DocumentError(f"Unsupported file type {ext!r}. Upload a .pdf, .docx, or .doc file.", 400)

    data = f.read()
    if not data:
        raise DocumentError("Uploaded file is empty", 400)

    if norm_type == "pdf":
        try:
            with fitz.open(stream=data, filetype="pdf") as d:
                if d.page_count == 0:
                    raise DocumentError("PDF has no pages", 400)
        except DocumentError:
            raise
        except Exception as exc:
            raise DocumentError(f"File is not a valid PDF: {exc}", 400) from exc
    else:
        try:
            DocxDocument(io.BytesIO(data))
        except Exception as exc:
            raise DocumentError(f"File is not a valid Word document: {exc}", 400) from exc

    stem = slugify_report_name(Path(f.filename).stem)
    doc_id = stem
    while any((DOCUMENTS_DIR / f"{doc_id}{e}").exists() for e in (".pdf", ".docx", ".doc")):
        doc_id = f"{stem}-{uuid.uuid4().hex[:6]}"

    out_path = DOCUMENTS_DIR / f"{doc_id}{ext}"
    out_path.write_bytes(data)

    return jsonify({"id": doc_id, "type": norm_type, "filename": out_path.name})


# ---------------------------------------------------------------------------
# API: document info / rendering
# ---------------------------------------------------------------------------

@app.route("/api/doc/<doc_id>/info")
def api_doc_info(doc_id):
    path, _ = resolve_pdf_path(doc_id, request.args.get("type", "pdf"))
    with fitz.open(path) as d:
        pages = [{"width": p.rect.width, "height": p.rect.height} for p in d]
    return jsonify({"page_count": len(pages), "pages": pages})


@app.route("/api/doc/<doc_id>/render/<int:page>")
def api_render(doc_id, page):
    path, _ = resolve_pdf_path(doc_id, request.args.get("type", "pdf"))
    try:
        dpi = int(request.args.get("dpi", 150))
    except ValueError:
        dpi = 150
    dpi = max(50, min(dpi, 600))

    with fitz.open(path) as d:
        if not (1 <= page <= d.page_count):
            raise DocumentError("Page out of range", 404)
        zoom = dpi / 72
        pix = d[page - 1].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        png_bytes = pix.tobytes("png")

    resp = Response(png_bytes, mimetype="image/png")
    resp.headers["Cache-Control"] = "no-store"
    return resp


# ---------------------------------------------------------------------------
# API: annotations (rectangles + freehand strokes)
# ---------------------------------------------------------------------------

@app.route("/api/doc/<doc_id>/annotations/<int:page>", methods=["GET", "POST"])
def api_annotations(doc_id, page):
    check_doc_id(doc_id)
    norm_type = normalize_type(request.args.get("type", "pdf"))
    path = annotations_path(doc_id, norm_type)

    if request.method == "GET":
        data = load_json(path, default={})
        return jsonify(data.get(str(page), []))

    body = request.get_json(silent=True) or {}
    anns = body.get("annotations")
    if not isinstance(anns, list):
        raise DocumentError("Body must contain an 'annotations' list", 400)

    data = load_json(path, default={})
    data[str(page)] = anns
    save_json(path, data)
    return jsonify({"status": "ok", "count": len(anns)})


# ---------------------------------------------------------------------------
# API: rectangular snippet extraction
# ---------------------------------------------------------------------------

@app.route("/api/doc/<doc_id>/snippet/<int:page>", methods=["POST"])
def api_create_snippet(doc_id, page):
    raw_type = request.args.get("type", "pdf")
    pdf_path, norm_type = resolve_pdf_path(doc_id, raw_type)

    body = request.get_json(silent=True) or {}
    try:
        x, y, w, h = float(body["x"]), float(body["y"]), float(body["w"]), float(body["h"])
    except (KeyError, TypeError, ValueError):
        raise DocumentError("Body must contain numeric x, y, w, h fractions (0-1)", 400)
    if w <= 0 or h <= 0 or not (0 <= x <= 1 and 0 <= y <= 1):
        raise DocumentError("Invalid rectangle: x, y must be in [0,1] and w, h > 0", 400)

    try:
        dpi = int(body.get("dpi", 300))
    except (TypeError, ValueError):
        dpi = 300
    dpi = max(72, min(dpi, 900))

    annotations = sanitize_annotations(body.get("annotations"))

    with fitz.open(pdf_path) as d:
        if not (1 <= page <= d.page_count):
            raise DocumentError("Page out of range", 404)
        p = d[page - 1]
        pr = p.rect
        draw_annotations_on_page(p, annotations, pr)
        clip = fitz.Rect(
            pr.x0 + x * pr.width,
            pr.y0 + y * pr.height,
            pr.x0 + min(x + w, 1.0) * pr.width,
            pr.y0 + min(y + h, 1.0) * pr.height,
        ) & pr
        zoom = dpi / 72
        pix = p.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=clip)
        png_bytes = pix.tobytes("png")

    snippet_id = uuid.uuid4().hex[:12]
    out_dir = snippets_dir(doc_id, norm_type)
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"p{page}_{snippet_id}.png"
    (out_dir / filename).write_bytes(png_bytes)

    meta_path = snippets_meta_path(doc_id, norm_type)
    meta = load_json(meta_path, default=[])
    entry = {
        "id": snippet_id,
        "page": page,
        "filename": filename,
        "rect": {"x": x, "y": y, "w": w, "h": h},
        "annotated": bool(annotations),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    meta.append(entry)
    save_json(meta_path, meta)

    result = dict(entry)
    result["url"] = url_for("api_snippet_file", doc_id=doc_id, filename=filename, type=raw_type)
    return jsonify(result)


@app.route("/api/doc/<doc_id>/download")
def api_download_annotated(doc_id):
    raw_type = request.args.get("type", "pdf")
    pdf_path, norm_type = resolve_pdf_path(doc_id, raw_type)

    data = load_json(annotations_path(doc_id, norm_type), default={})

    with fitz.open(pdf_path) as d:
        for page_key, raw_anns in data.items():
            try:
                page_num = int(page_key)
            except ValueError:
                continue
            if not (1 <= page_num <= d.page_count):
                continue
            annotations = sanitize_annotations(raw_anns)
            if not annotations:
                continue
            p = d[page_num - 1]
            draw_annotations_on_page(p, annotations, p.rect)
        pdf_bytes = d.tobytes(deflate=True)

    resp = Response(pdf_bytes, mimetype="application/pdf")
    resp.headers["Cache-Control"] = "no-store"
    resp.headers["Content-Disposition"] = f'attachment; filename="{doc_id}-annotated.pdf"'
    return resp


@app.route("/api/doc/<doc_id>/snippets")
def api_list_snippets(doc_id):
    check_doc_id(doc_id)
    norm_type = normalize_type(request.args.get("type", "pdf"))
    page = request.args.get("page", type=int)

    meta = load_json(snippets_meta_path(doc_id, norm_type), default=[])
    if page is not None:
        meta = [m for m in meta if m["page"] == page]
    raw_type = request.args.get("type", "pdf")
    for m in meta:
        m["url"] = url_for("api_snippet_file", doc_id=doc_id, filename=m["filename"], type=raw_type)
    return jsonify(meta)


@app.route("/api/doc/<doc_id>/snippet/<snippet_id>", methods=["DELETE"])
def api_delete_snippet(doc_id, snippet_id):
    check_doc_id(doc_id)
    norm_type = normalize_type(request.args.get("type", "pdf"))
    meta_path = snippets_meta_path(doc_id, norm_type)
    meta = load_json(meta_path, default=[])
    remaining = [m for m in meta if m["id"] != snippet_id]
    removed = [m for m in meta if m["id"] == snippet_id]
    if not removed:
        raise DocumentError(f"No snippet with id {snippet_id}", 404)
    save_json(meta_path, remaining)
    for m in removed:
        f = snippets_dir(doc_id, norm_type) / m["filename"]
        if f.exists():
            f.unlink()
    return jsonify({"status": "ok"})


@app.route("/api/reports", methods=["GET", "POST"])
def api_reports():
    if request.method == "GET":
        items = []
        for f in REPORTS_DIR.glob("*.json"):
            data = load_json(f, default=None)
            if not isinstance(data, dict):
                continue
            items.append({
                "id": f.stem,
                "name": data.get("name") or f.stem,
                "source_doc": data.get("source_doc", ""),
                "source_type": data.get("source_type", "pdf"),
                "created_at": data.get("created_at", ""),
                "updated_at": data.get("updated_at", ""),
            })
        items.sort(key=lambda x: x["updated_at"], reverse=True)
        return jsonify(items)

    body = request.get_json(silent=True) or {}
    name = str(body.get("name") or "").strip()[:200]
    if not name:
        raise DocumentError("A document name is required", 400)

    source_doc = str(body.get("source_doc") or "")
    source_type = "pdf"
    if source_doc:
        check_doc_id(source_doc)
        source_type = normalize_type(body.get("source_type", "pdf"))

    report_id = f"{slugify_report_name(name)}-{uuid.uuid4().hex[:6]}"
    now = datetime.now(timezone.utc).isoformat()
    data = {
        "name": name,
        "html": "",
        "source_doc": source_doc,
        "source_type": source_type,
        "margins": dict(REPORT_DEFAULT_MARGINS),
        "created_at": now,
        "updated_at": now,
    }
    save_json(report_path(report_id), data)
    return jsonify({"id": report_id, **data})


@app.route("/api/report/<report_id>", methods=["GET", "POST"])
def api_report(report_id):
    check_report_id(report_id)
    path = report_path(report_id)
    existing = load_json(path, default=None)
    if existing is None:
        raise DocumentError(f"No document with id {report_id!r}", 404)

    if request.method == "GET":
        resp = dict(existing)
        resp["margins"] = sanitize_margins(existing.get("margins"))
        return jsonify(resp)

    body = request.get_json(silent=True) or {}
    name = str(body.get("name", existing.get("name", ""))).strip()[:200]
    if not name:
        raise DocumentError("A document name is required", 400)

    source_doc = str(body.get("source_doc", existing.get("source_doc", "")))
    source_type = "pdf"
    if source_doc:
        check_doc_id(source_doc)
        source_type = normalize_type(body.get("source_type", existing.get("source_type", "pdf")))

    html_content = sanitize_report_html(body.get("html", existing.get("html", "")))
    margins = sanitize_margins(body.get("margins"), existing.get("margins"))
    data = {
        "name": name,
        "html": html_content,
        "source_doc": source_doc,
        "source_type": source_type,
        "margins": margins,
        "created_at": existing.get("created_at", datetime.now(timezone.utc).isoformat()),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    save_json(path, data)
    return jsonify(data)


@app.route("/api/report/<report_id>/export")
def api_report_export(report_id):
    check_report_id(report_id)
    data = load_json(report_path(report_id), default=None)
    if data is None:
        raise DocumentError(f"No document with id {report_id!r}", 404)

    title = data.get("name") or report_id
    body_html = inline_report_images(data.get("html") or "")
    if not body_html.strip():
        raise DocumentError("Document is empty — add some content before exporting", 400)

    pdf_bytes = render_report_pdf(title, body_html, data.get("margins"))

    resp = Response(pdf_bytes, mimetype="application/pdf")
    resp.headers["Cache-Control"] = "no-store"
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "-", title).strip("-") or report_id
    resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}.pdf"'
    return resp


@app.route("/api/report/<report_id>/export.docx")
def api_report_export_docx(report_id):
    check_report_id(report_id)
    data = load_json(report_path(report_id), default=None)
    if data is None:
        raise DocumentError(f"No document with id {report_id!r}", 404)

    title = data.get("name") or report_id
    body_html = inline_report_images(data.get("html") or "")
    if not body_html.strip():
        raise DocumentError("Document is empty — add some content before exporting", 400)

    docx_bytes = render_report_docx(title, body_html, data.get("margins"))

    resp = Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp.headers["Cache-Control"] = "no-store"
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "-", title).strip("-") or report_id
    resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}.docx"'
    return resp


@app.route("/media/snippets/<doc_id>/<path:filename>")
def api_snippet_file(doc_id, filename):
    check_doc_id(doc_id)
    norm_type = normalize_type(request.args.get("type", "pdf"))
    directory = snippets_dir(doc_id, norm_type)
    if "/" in filename or "\\" in filename:
        abort(400)
    return send_from_directory(directory, filename)


if __name__ == "__main__":
    import os

    port = int(os.environ.get("PORT", 5050))
    app.run(host="127.0.0.1", port=port, debug=True)
